import pdfplumber
from docx import Document
import os

def extract_text_from_pdf(file_path):
    text = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
    return "\n".join(text)

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    parts = []

    # 1) Main body paragraphs
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)

    # 2) Tables (each cell's text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    t = para.text.strip()
                    if t:
                        parts.append(t)

    # 3) Headers and footers across all sections
    for section in doc.sections:
        for hdr in [section.header, section.first_page_header, section.even_page_header]:
            if hdr is not None:
                for para in hdr.paragraphs:
                    t = para.text.strip()
                    if t:
                        parts.append(t)
        for ftr in [section.footer, section.first_page_footer, section.even_page_footer]:
            if ftr is not None:
                for para in ftr.paragraphs:
                    t = para.text.strip()
                    if t:
                        parts.append(t)

    return "\n".join(parts)

def extract_text_from_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file format")